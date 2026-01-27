"""
Document Builder Service

Handles the generation of Word documents using python-docx.
Creates properly formatted HB Weekly Reports with styles, tables, and metadata.
"""

import logging
from dataclasses import dataclass, field
from datetime import datetime, date
from pathlib import Path
from typing import Dict, List, Optional, Any

from ..config.settings import HBWeeklyReportConfig, OutputConfig
from ..agents.report_writer_agent import ReportContent

logger = logging.getLogger(__name__)


@dataclass
class DocumentResult:
    """Result of document generation"""
    success: bool
    document_path: Optional[Path] = None
    file_size_bytes: int = 0
    generation_time_seconds: float = 0.0
    error_message: Optional[str] = None


class DocumentBuilder:
    """
    Builds Word documents from ReportContent

    Features:
    - Styled headings and body text
    - Formatted price tables
    - Spread tables with % of full carry
    - Proper document metadata
    - Template support (optional)
    """

    def __init__(self, config: HBWeeklyReportConfig):
        """
        Initialize Document Builder

        Args:
            config: HB Weekly Report configuration
        """
        self.config = config
        self.output_config = config.output
        self.logger = logging.getLogger(f"{self.__class__.__name__}")

        # Verify python-docx is available
        self._docx_available = self._check_docx_available()

        self.logger.info(f"Initialized DocumentBuilder, docx available: {self._docx_available}")

    def _check_docx_available(self) -> bool:
        """Check if python-docx is available"""
        try:
            import docx
            return True
        except ImportError:
            self.logger.warning("python-docx not installed. Install with: pip install python-docx")
            return False

    def build_document(self, content: ReportContent) -> DocumentResult:
        """
        Build Word document from report content

        Args:
            content: ReportContent with all sections

        Returns:
            DocumentResult with document path
        """
        start_time = datetime.utcnow()

        if not self._docx_available:
            return DocumentResult(
                success=False,
                error_message="python-docx not installed"
            )

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            # Create document
            if self.output_config.template_file:
                template_path = self.output_config.template_directory / self.output_config.template_file
                if template_path.exists():
                    doc = Document(str(template_path))
                else:
                    doc = Document()
            else:
                doc = Document()

            # Set document properties
            self._set_document_properties(doc, content)

            # Add title
            self._add_title(doc, content)

            # Add Executive Summary
            self._add_section(doc, "Executive Summary", content.executive_summary, level=1)

            # Add Macro and Weather Update
            self._add_section(doc, "Macro and Weather Update", "", level=1)
            self._add_section(doc, "Macro Developments", content.macro_update, level=2)
            self._add_section(doc, "Weather Update", content.weather_update, level=2)

            # Add Commodity Deep Dives
            self._add_section(doc, "Market Analysis", "", level=1)

            commodity_titles = {
                "corn": "Corn",
                "wheat": "Wheat",
                "soybeans": "Soybeans",
                "soybean_meal": "Soybean Meal",
                "soybean_oil": "Soybean Oil",
            }

            for commodity, title in commodity_titles.items():
                if commodity in content.commodity_sections:
                    self._add_section(doc, title, content.commodity_sections[commodity], level=2)

            # Add Price Tables
            self._add_section(doc, "Weekly Prices and Spreads", "", level=1)

            if content.price_table_data.get("futures"):
                self._add_price_table(doc, "Futures Prices", content.price_table_data["futures"])

            if content.spread_table_data:
                self._add_spread_table(doc, "Key Spreads", content.spread_table_data)

            if content.international_table_data:
                self._add_price_table(doc, "International Prices", content.international_table_data)

            # Add Synthesis and Outlook
            self._add_section(doc, "Synthesis and Outlook", content.synthesis_outlook, level=1)

            # Add Key Triggers
            self._add_key_triggers(doc, content.key_triggers)

            # Add metadata note
            self._add_metadata_note(doc, content)

            # Generate filename
            filename = self._generate_filename(content.report_date)
            output_dir = Path(self.output_config.output_directory)
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / filename

            # Save document
            doc.save(str(output_path))

            end_time = datetime.utcnow()
            duration = (end_time - start_time).total_seconds()

            self.logger.info(f"Document saved to {output_path}")

            return DocumentResult(
                success=True,
                document_path=output_path,
                file_size_bytes=output_path.stat().st_size,
                generation_time_seconds=duration,
            )

        except Exception as e:
            self.logger.error(f"Document generation failed: {e}", exc_info=True)
            return DocumentResult(
                success=False,
                error_message=str(e),
            )

    def _set_document_properties(self, doc, content: ReportContent):
        """Set document metadata properties"""
        try:
            core_props = doc.core_properties
            core_props.author = self.output_config.author
            core_props.title = f"HigbyBarrett Weekly Report - {content.report_date.strftime('%B %d, %Y')}"
            core_props.subject = "Agricultural Commodity Market Analysis"
            core_props.keywords = "corn, wheat, soybeans, commodities, market analysis"
            core_props.comments = f"Generated by {self.config.agent_name} v{self.config.agent_version}"
            core_props.created = datetime.utcnow()
        except Exception as e:
            self.logger.warning(f"Could not set document properties: {e}")

    def _add_title(self, doc, content: ReportContent):
        """Add document title"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title = doc.add_heading(level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.add_run("HigbyBarrett Weekly Report")
        run.font.size = Pt(24)

        # Add date subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"Week Ending {content.week_ending.strftime('%B %d, %Y')}")

        doc.add_paragraph()  # Spacing

    def _add_section(self, doc, title: str, content: str, level: int = 1):
        """Add a section with heading and content"""
        from docx.shared import Pt

        # Add heading
        heading = doc.add_heading(title, level=level)

        # Add content if provided
        if content:
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    if para_text.startswith("- "):
                        # Bullet point
                        for line in para_text.split("\n"):
                            if line.strip().startswith("- "):
                                bullet_text = line.strip()[2:]
                                doc.add_paragraph(bullet_text, style='List Bullet')
                    else:
                        para = doc.add_paragraph(para_text)
                        para.style = 'Normal'

    def _add_price_table(self, doc, title: str, data: List[Dict]):
        """Add a formatted price table"""
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        if not data:
            return

        doc.add_heading(title, level=3)

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Contract', 'Current', 'Week Ago', 'Year Ago', 'Week Chg']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for row_data in data:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data.get('name', '')[:30]  # Truncate long names

            current = row_data.get('current')
            week_ago = row_data.get('week_ago')
            year_ago = row_data.get('year_ago')
            week_change = row_data.get('week_change')

            row_cells[1].text = f"{current:.2f}" if current is not None else "N/A"
            row_cells[2].text = f"{week_ago:.2f}" if week_ago is not None else "N/A"
            row_cells[3].text = f"{year_ago:.2f}" if year_ago is not None else "N/A"

            if week_change is not None:
                change_text = f"{week_change:+.2f}"
                row_cells[4].text = change_text
            else:
                row_cells[4].text = "N/A"

        # Add table note
        note = doc.add_paragraph()
        note.add_run("Source: USDA AMS, CME Group. Prices as of report date.").italic = True

        doc.add_paragraph()  # Spacing

    def _add_spread_table(self, doc, title: str, data: List[Dict]):
        """Add spread table with % of full carry"""
        from docx.shared import Pt

        if not data:
            return

        doc.add_heading(title, level=3)

        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Spread', 'Value (cents)', 'Week Chg', '% Full Carry']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for row_data in data:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data.get('name', '')

            current = row_data.get('current')
            week_change = row_data.get('week_change')
            pct_carry = row_data.get('pct_full_carry')

            row_cells[1].text = f"{current:.2f}" if current is not None else "N/A"
            row_cells[2].text = f"{week_change:+.2f}" if week_change is not None else "N/A"
            row_cells[3].text = f"{pct_carry:.1f}%" if pct_carry is not None else "N/A"

        # Add footnote
        note = doc.add_paragraph()
        storage_cost = self.config.commodities.storage_cost_per_month
        interest_rate = self.config.commodities.interest_rate_annual * 100
        note.add_run(
            f"Full carry assumes ${storage_cost}/bu/month storage + {interest_rate:.0f}% interest."
        ).italic = True

        doc.add_paragraph()

    def _add_key_triggers(self, doc, triggers: List[str]):
        """Add key triggers watchlist section"""
        doc.add_heading("Key Fundamental Triggers to Watch", level=1)

        if triggers:
            for trigger in triggers:
                doc.add_paragraph(trigger, style='List Bullet')
        else:
            doc.add_paragraph("No significant triggers identified for the coming week.")

    def _add_metadata_note(self, doc, content: ReportContent):
        """Add metadata/footnote section"""
        from docx.shared import Pt, RGBColor

        doc.add_page_break()

        # Metadata section
        doc.add_heading("Report Metadata", level=2)

        metadata_items = [
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
            f"Agent: {self.config.agent_name} v{self.config.agent_version}",
            f"Data Sources: {', '.join(content.data_sources)}",
        ]

        if content.placeholders:
            metadata_items.append(f"Placeholders: {len(content.placeholders)}")

        if content.llm_estimates:
            metadata_items.append(f"LLM-Derived Estimates: {len(content.llm_estimates)}")
            for estimate in content.llm_estimates[:5]:
                metadata_items.append(f"  - {estimate}")

        for item in metadata_items:
            para = doc.add_paragraph(item)
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)

    def _generate_filename(self, report_date: date) -> str:
        """Generate filename for the report"""
        pattern = self.output_config.filename_pattern
        filename = pattern.format(
            month=report_date.strftime('%B'),
            day=report_date.day,
            year=report_date.year,
        )
        return f"{filename}.docx"


class PlaceholderDocumentBuilder(DocumentBuilder):
    """
    Document builder that creates placeholder documents when python-docx is unavailable
    Outputs to HTML or plain text as fallback
    """

    def build_document(self, content: ReportContent) -> DocumentResult:
        """Build fallback document format"""
        start_time = datetime.utcnow()

        try:
            # Generate as HTML
            html_content = self._generate_html(content)

            # Save as HTML
            filename = self._generate_filename(content.report_date).replace('.docx', '.html')
            output_dir = Path(self.output_config.output_directory)
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / filename

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            end_time = datetime.utcnow()

            return DocumentResult(
                success=True,
                document_path=output_path,
                file_size_bytes=output_path.stat().st_size,
                generation_time_seconds=(end_time - start_time).total_seconds(),
            )

        except Exception as e:
            return DocumentResult(
                success=False,
                error_message=str(e),
            )

    def _generate_html(self, content: ReportContent) -> str:
        """Generate HTML version of report"""
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>HB Weekly Report - {content.report_date}</title>",
            "<style>",
            "body { font-family: Calibri, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
            "h1 { color: #2c3e50; }",
            "h2 { color: #34495e; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #4472C4; color: white; }",
            ".metadata { color: #888; font-size: 0.9em; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>HigbyBarrett Weekly Report</h1>",
            f"<p>Week Ending {content.week_ending.strftime('%B %d, %Y')}</p>",
            "<hr>",
            "<h2>Executive Summary</h2>",
            f"<p>{content.executive_summary}</p>",
            "<h2>Macro and Weather Update</h2>",
            f"<p>{content.macro_update}</p>",
            f"<p>{content.weather_update}</p>",
        ]

        # Add commodity sections
        for commodity, section_content in content.commodity_sections.items():
            html_parts.append(f"<h2>{commodity.title()}</h2>")
            html_parts.append(f"<p>{section_content}</p>")

        # Add synthesis
        html_parts.append("<h2>Synthesis and Outlook</h2>")
        html_parts.append(f"<p>{content.synthesis_outlook}</p>")

        # Add triggers
        html_parts.append("<h2>Key Triggers</h2>")
        html_parts.append("<ul>")
        for trigger in content.key_triggers:
            html_parts.append(f"<li>{trigger}</li>")
        html_parts.append("</ul>")

        # Metadata
        html_parts.append("<hr>")
        html_parts.append("<div class='metadata'>")
        html_parts.append(f"<p>Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</p>")
        html_parts.append(f"<p>Data Sources: {', '.join(content.data_sources)}</p>")
        html_parts.append("</div>")

        html_parts.append("</body>")
        html_parts.append("</html>")

        return "\n".join(html_parts)
