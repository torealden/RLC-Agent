"""Render the markdown leave-behind to docx and copy to the Dropbox folder.

The original `RLC_Helios_LeaveBehind.docx` is backed up first.
"""
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

SRC = Path(r"C:\dev\RLC-Agent\docs\specs\helios_leave_behind.md")
DST_DIR = Path(r"C:\Users\torem\RLC Dropbox\Tore Alden\Misc Personal Stuff\Helios")
DST = DST_DIR / "RLC_Helios_LeaveBehind.docx"
BACKUP = DST_DIR / f"RLC_Helios_LeaveBehind_backup_{datetime.now():%Y%m%d_%H%M%S}.docx"

# Back up the original
if DST.exists():
    shutil.copy2(DST, BACKUP)
    print(f"Backed up to: {BACKUP.name}")

# Parse markdown into doc
md = SRC.read_text(encoding="utf-8")
doc = Document()

# Page setup: tight margins, 1.5pp goal
for section in doc.sections:
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

# Set default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

lines = md.splitlines()
in_list = False
for raw in lines:
    line = raw.rstrip()

    if not line:
        in_list = False
        continue

    if line.startswith("# "):
        # Title
        p = doc.add_paragraph()
        run = p.add_run(line[2:].strip())
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(2)
    elif line.startswith("## "):
        # Section heading
        p = doc.add_paragraph()
        run = p.add_run(line[3:].strip())
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Georgia'
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    elif line.startswith("**") and line.endswith("**"):
        # Subtitle / bold-only line
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*"))
        run.italic = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(8)
    elif line.startswith("---"):
        # Horizontal rule — render as small blank space
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    elif line.startswith("- "):
        # Bullet
        bullet_text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        # Handle inline **bold** in bullet text
        chunks = bullet_text.split("**")
        for i, chunk in enumerate(chunks):
            run = p.add_run(chunk)
            run.bold = (i % 2 == 1)
        p.paragraph_format.space_after = Pt(2)
        in_list = True
    elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
        # Italic-only line (signature line at the end)
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*"))
        run.italic = True
        run.font.size = Pt(9.5)
        p.paragraph_format.space_before = Pt(6)
    else:
        # Plain paragraph — handle inline **bold**
        p = doc.add_paragraph()
        chunks = line.split("**")
        for i, chunk in enumerate(chunks):
            run = p.add_run(chunk)
            run.bold = (i % 2 == 1)
        p.paragraph_format.space_after = Pt(4)

doc.save(str(DST))
print(f"Wrote: {DST}")
