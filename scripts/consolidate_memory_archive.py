"""Consolidate the agent memory register into durable artifacts.

The memory lives in ~/.claude/.../memory/ (143 topic files + MEMORY.md index), which is
OUTSIDE git and whose index has outgrown the session load limit (content truncates). This
writes the full content into two durable, single-file artifacts:

  docs/memory_archive/RLC_AGENT_MEMORY_ARCHIVE.md   (git-versioned = truly unlosable)
  docs/memory_archive/RLC_AGENT_MEMORY_ARCHIVE.docx (portable Word copy)

Secrets are REDACTED before writing (these artifacts are git-committed / portable): known
passwords + any 'password/secret/token = value' lines. The live memory files are untouched.

Run: python scripts/consolidate_memory_archive.py
"""
import re
import sys
from pathlib import Path

ROOT = Path(r"C:/dev/RLC-Agent")
MEM = Path.home() / ".claude" / "projects" / "C--dev-RLC-Agent" / "memory"
OUT = ROOT / "docs" / "memory_archive"

# redaction: known secret + generic secret-assignment lines
KNOWN_SECRETS = ["SoupBoss1"]
SECRET_LINE = re.compile(r'(?i)(password|passwd|secret|api[_-]?key|token)\s*[:=]\s*\S+')


def redact(text: str) -> tuple[str, int]:
    n = 0
    for s in KNOWN_SECRETS:
        if s in text:
            n += text.count(s); text = text.replace(s, "[REDACTED-SECRET]")
    out_lines = []
    for line in text.splitlines():
        if SECRET_LINE.search(line):
            line = SECRET_LINE.sub(r"\1: [REDACTED-SECRET]", line); n += 1
        out_lines.append(line)
    return "\n".join(out_lines), n


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    files = sorted([p for p in MEM.glob("*.md") if p.name != "MEMORY.md"])
    index = MEM / "MEMORY.md"

    total_redactions = 0
    sections = []  # (title, filename, body)
    # index first, then topic files grouped by type prefix
    order = [index] + files
    for p in order:
        raw = p.read_text(encoding="utf-8", errors="replace")
        body, n = redact(raw)
        total_redactions += n
        sections.append((p.stem, p.name, body))

    # ---- markdown archive ----
    from datetime import datetime  # NOTE: only for the archive header stamp
    md = ["# RLC-Agent — Consolidated Memory Archive", "",
          f"Durable, git-versioned snapshot of the agent memory register "
          f"({len(order)} files). Secrets redacted. Regenerate with "
          f"`python scripts/consolidate_memory_archive.py`.", "",
          "## Contents", ""]
    for title, fname, _ in sections:
        anchor = re.sub(r'[^a-z0-9]+', '-', title.lower()).strip('-')
        md.append(f"- [{title}](#{anchor})")
    md.append("\n---\n")
    for title, fname, body in sections:
        md.append(f"## {title}\n\n*(`{fname}`)*\n\n{body}\n\n---\n")
    md_text = "\n".join(md)
    (OUT / "RLC_AGENT_MEMORY_ARCHIVE.md").write_text(md_text, encoding="utf-8")

    # ---- docx archive ----
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading("RLC-Agent — Consolidated Memory Archive", level=0)
    doc.add_paragraph(f"Durable snapshot of the agent memory register ({len(order)} files). "
                      "Secrets redacted.")
    for title, fname, body in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(fname).italic = True
        for block in body.split("\n"):
            b = block.rstrip()
            if not b:
                continue
            if b.startswith("#"):
                lvl = min(b.count("#"), 4) + 1
                doc.add_heading(b.lstrip("# ").strip(), level=min(lvl, 4))
            else:
                doc.add_paragraph(b)
    doc.save(str(OUT / "RLC_AGENT_MEMORY_ARCHIVE.docx"))

    md_kb = (OUT / "RLC_AGENT_MEMORY_ARCHIVE.md").stat().st_size / 1024
    docx_kb = (OUT / "RLC_AGENT_MEMORY_ARCHIVE.docx").stat().st_size / 1024
    print(f"Consolidated {len(order)} memory files:")
    print(f"  {OUT / 'RLC_AGENT_MEMORY_ARCHIVE.md'}  ({md_kb:.0f} KB)")
    print(f"  {OUT / 'RLC_AGENT_MEMORY_ARCHIVE.docx'} ({docx_kb:.0f} KB)")
    print(f"  {total_redactions} secret(s) redacted")


if __name__ == "__main__":
    main()
