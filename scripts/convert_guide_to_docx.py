"""
Convert AUTONOMOUS_PIPELINE_GUIDE.md to a styled Word document.
Uses python-docx to create a professional .docx matching RLC document style.

Usage:
    python scripts/convert_guide_to_docx.py
"""

import re
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def create_styles(doc):
    """Set up custom styles for the document."""
    # Modify default Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level in range(1, 5):
        style_name = f'Heading {level}'
        heading_style = doc.styles[style_name]
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
        if level == 1:
            heading_style.font.size = Pt(20)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        elif level == 3:
            heading_style.font.size = Pt(13)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        elif level == 4:
            heading_style.font.size = Pt(11)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(10)
            heading_style.paragraph_format.space_after = Pt(4)

    # Code block style
    try:
        code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code_style = doc.styles['CodeBlock']
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Inches(0.3)


def add_title_page(doc):
    """Add a styled title page."""
    # Add some spacing
    for _ in range(4):
        doc.add_paragraph('')

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Autonomous Pipeline Architecture Guide')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.name = 'Calibri'

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('for Round Lakes Companies')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.name = 'Calibri'

    doc.add_paragraph('')

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('From Data Collection to Walk-Away Report Generation')
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add spacing
    for _ in range(6):
        doc.add_paragraph('')

    # Version info
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f'Version 1.0 — March 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break
    doc.add_page_break()


def parse_markdown_table(lines):
    """Parse a markdown table into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator rows
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(cell_text.replace('**', ''))
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph('')  # Space after table


def add_code_block(doc, code_lines):
    """Add a code block to the document."""
    for line in code_lines:
        p = doc.add_paragraph(style='CodeBlock')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


def add_checkbox_item(doc, text, checked=False):
    """Add a checkbox item."""
    p = doc.add_paragraph()
    prefix = '[X] ' if checked else '[ ] '
    run = p.add_run(prefix + text.replace('**', ''))
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)


def process_inline_formatting(paragraph, text):
    """Add text to a paragraph with bold/italic/code formatting."""
    # Simple approach: split on ** for bold, * for italic, ` for code
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)


def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to styled Word document."""
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    create_styles(doc)
    add_title_page(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    # Skip the title and front matter (already handled by title page)
    # Find where Section 1 starts
    while i < len(lines):
        if lines[i].strip().startswith('## 1.'):
            break
        i += 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    add_table(doc, parse_markdown_table(table_lines))
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            add_table(doc, parse_markdown_table(table_lines))
            table_lines = []
            in_table = False

        # Skip TOC and horizontal rules
        if stripped == '---' or stripped.startswith('[') and '](#' in stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('####'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(text, level=4)
        elif stripped.startswith('###'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(text, level=3)
        elif stripped.startswith('## '):
            text = stripped.lstrip('#').strip()
            # Page break before major sections (## N.)
            if re.match(r'^\d+\.', text):
                doc.add_page_break()
            doc.add_heading(text, level=2)
        elif stripped.startswith('# '):
            # Skip top-level title (handled by title page)
            pass

        # Checkbox items
        elif stripped.startswith('- [ ]') or stripped.startswith('- [x]') or stripped.startswith('- [X]'):
            checked = '[x]' in stripped.lower()
            text = re.sub(r'^- \[.\]\s*', '', stripped)
            add_checkbox_item(doc, text, checked)

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)

        # Numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)

        # Empty lines
        elif not stripped:
            pass

        # Regular paragraphs
        else:
            p = doc.add_paragraph()
            process_inline_formatting(p, stripped)

        i += 1

    # Flush any remaining table
    if in_table:
        add_table(doc, parse_markdown_table(table_lines))

    doc.save(str(docx_path))
    print(f"[OK] Document saved: {docx_path}")
    print(f"     Size: {Path(docx_path).stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    md_path = Path(r'C:\dev\RLC-Agent\docs\AUTONOMOUS_PIPELINE_GUIDE.md')

    # Primary output: alongside the project docs
    docx_path = Path(r'C:\dev\RLC-Agent\docs\Autonomous_Pipeline_Architecture_Guide_v1.1.docx')

    # Secondary output: Dropbox for sharing
    dropbox_path = Path(r'C:\Users\torem\Dropbox\AI Stuff\Desktop LLM Project\Autonomous_Pipeline_Architecture_Guide.docx')

    convert_md_to_docx(md_path, docx_path)

    # Copy to Dropbox
    if dropbox_path.parent.exists():
        convert_md_to_docx(md_path, dropbox_path)
        print(f"[OK] Dropbox copy saved: {dropbox_path}")
    else:
        print(f"[WARN] Dropbox path not found: {dropbox_path.parent}")
