"""Generate analyst_database_setup_guide.docx for Felipe onboarding."""
from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('RLC Commodities Database', level=0)
doc.add_heading('Analyst Setup Guide', level=1)

# Overview
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The RLC Commodities database is a PostgreSQL database hosted on AWS (Amazon RDS). '
    'It contains agricultural commodity data organized in three layers:'
)
doc.add_paragraph('Bronze \u2014 Raw data as collected from sources (USDA, EIA, CFTC, etc.)', style='List Bullet')
doc.add_paragraph('Silver \u2014 Cleaned and standardized data with calculated fields', style='List Bullet')
doc.add_paragraph('Gold \u2014 Analytics-ready views, balance sheets, and matrix views for spreadsheets', style='List Bullet')
doc.add_paragraph(
    'You can connect from: pgAdmin 4 (free database GUI), Excel VBA macros '
    '(one-click spreadsheet updates), or DBeaver.'
)

# Connection details
doc.add_heading('Database Connection Details', level=2)
table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('Setting', 'Value'),
    ('Host', 'rlc-commodities.c16c6wm826t7.us-east-2.rds.amazonaws.com'),
    ('Port', '5432'),
    ('Database', 'rlc_commodities'),
    ('Username', 'postgres'),
    ('Password', 'SoupBoss1'),
    ('SSL Mode', 'Require'),
]):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    if i == 0:
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

# Step 1: pgAdmin
doc.add_heading('Step 1: Install pgAdmin 4', level=2)
doc.add_paragraph('pgAdmin is the standard PostgreSQL GUI for browsing tables and running queries.')
for step in [
    '1. Go to https://www.pgadmin.org/download/pgadmin-4-windows/',
    '2. Download the latest Windows installer (.exe)',
    '3. Run the installer with default settings',
    '4. Set a master password when prompted (this is for pgAdmin itself)',
]:
    doc.add_paragraph(step)

doc.add_heading('Connect to RLC Database in pgAdmin', level=3)
for step in [
    'Right-click Servers in the left panel > Register > Server',
    'General tab: Name = "RLC Commodities"',
    'Connection tab: enter Host, Port, Database, Username, Password from table above',
    'Check "Save password"',
    'SSL tab: SSL mode = "Require"',
    'Click Save',
]:
    doc.add_paragraph(step, style='List Bullet')
doc.add_paragraph(
    'You should now see the database in the left panel. '
    'Expand it to browse schemas (bronze, silver, gold) and their tables/views.'
)

doc.add_heading('Running Queries', level=3)
doc.add_paragraph(
    'Click on rlc_commodities in the left panel, then click the Query Tool icon '
    '(or Tools > Query Tool). Type SQL and press F5 to execute.'
)

doc.add_heading('Example Queries', level=3)
queries = [
    ('US corn balance sheet (latest 3 years)',
     'SELECT * FROM gold.fas_us_corn_balance_sheet ORDER BY marketing_year DESC LIMIT 3;'),
    ('Current CFTC managed money positioning',
     'SELECT * FROM gold.cftc_sentiment;'),
    ('Weekly ethanol production',
     'SELECT * FROM gold.eia_ethanol_weekly ORDER BY week_ending DESC LIMIT 10;'),
    ('Brazil soybean production by state',
     "SELECT * FROM gold.brazil_soybean_production WHERE crop_year = '2024/25' ORDER BY production DESC;"),
    ('Monthly soybean inspections (thousand bushels)',
     "SELECT * FROM gold.fgis_inspections_monthly_matrix_kbu WHERE grain = 'SOYBEANS' AND year = 2025 ORDER BY month, spreadsheet_row;"),
]
for desc, sql in queries:
    p = doc.add_paragraph()
    run = p.add_run(desc + ': ')
    run.bold = True
    run = p.add_run(sql)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

# Step 2: ODBC
doc.add_heading('Step 2: Install PostgreSQL ODBC Driver (for Excel VBA)', level=2)
doc.add_paragraph('The Excel VBA macros need this driver to connect to the database.')
for step in [
    '1. Go to https://www.postgresql.org/ftp/odbc/versions/msi/',
    '2. Download the latest 64-bit installer (psqlodbc_xx_xx_xxxx-x64.zip)',
    '3. Extract the ZIP and run the .msi installer with defaults',
]:
    doc.add_paragraph(step)
doc.add_paragraph(
    'Verify: Windows search > "ODBC" > ODBC Data Sources (64-bit) > '
    'Drivers tab > look for "PostgreSQL UNICODE(x64)"'
)

# Step 3: ActiveX
doc.add_heading('Step 3: Enable ActiveX Data Objects in Excel', level=2)
doc.add_paragraph('Usually pre-installed with Office. To verify:')
for step in [
    '1. Open Excel > Alt+F11 (VBA editor)',
    '2. Tools > References',
    '3. Check "Microsoft ActiveX Data Objects 6.1 Library"',
    '4. Click OK',
]:
    doc.add_paragraph(step)

# Step 4: VBA
doc.add_heading('Step 4: Import VBA Modules into Workbooks', level=2)
doc.add_paragraph('Each workbook has VBA modules for one-click database updates:')
for step in [
    '1. Open the workbook in Excel',
    '2. Press Alt+F11 to open the VBA editor',
    '3. Right-click the workbook name in the left panel > Import File',
    '4. Navigate to the src/tools/ folder',
    '5. Select the appropriate .bas file and click Open',
    '6. Close VBA editor, save the workbook as .xlsm',
]:
    doc.add_paragraph(step)

doc.add_heading('VBA Module Reference', level=3)
vba_table = doc.add_table(rows=9, cols=4, style='Light Grid Accent 1')
headers = ['Module File', 'Shortcut', 'Workbook', 'Description']
for j, h in enumerate(headers):
    vba_table.rows[0].cells[j].text = h
    for p in vba_table.rows[0].cells[j].paragraphs:
        for r in p.runs:
            r.bold = True
vba_data = [
    ('TradeUpdaterSQL.bas', 'Ctrl+I', 'Census Trade', 'US import/export trade'),
    ('InspectionsUpdaterSQL.bas', 'Ctrl+G', 'FGIS Inspections', 'Export inspections (000 bu)'),
    ('CrushUpdaterSQL.bas', 'Ctrl+U', 'Crush Data', 'Soybean crush'),
    ('BiofuelDataUpdater.bas', 'Ctrl+B', 'Biofuel S&D', 'Biofuel balance sheets'),
    ('FeedstockUpdaterSQL.bas', 'Ctrl+E', 'US Feedstock', 'EIA ethanol + petroleum'),
    ('EMTSDataUpdater.bas', 'Ctrl+E', 'EMTS RIN', 'EPA RIN generation'),
    ('RINUpdaterSQL.bas', 'Ctrl+R', 'RIN Data', 'RIN transactions'),
    ('EIAFeedstockUpdater.bas', 'Ctrl+D', 'EIA Feedstock', 'EIA feedstock data'),
]
for i, row_data in enumerate(vba_data):
    for j, val in enumerate(row_data):
        vba_table.rows[i + 1].cells[j].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'Pattern: Ctrl+[letter] = quick update (latest N periods). '
    'Ctrl+Shift+[letter] = custom update (you choose how many periods).'
)

# Key views
doc.add_heading('Key Database Views for Analysts', level=2)

sections = [
    ('Balance Sheets', [
        ('gold.fas_us_corn_balance_sheet', 'US corn S&D'),
        ('gold.fas_us_soybeans_balance_sheet', 'US soybeans S&D'),
        ('gold.fas_us_wheat_balance_sheet', 'US wheat S&D'),
        ('gold.us_soybean_balance_sheet', 'Historical US soybean S&D (ERS)'),
        ('gold.us_soybean_oil_balance_sheet', 'US soybean oil S&D'),
        ('gold.us_soybean_meal_balance_sheet', 'US soybean meal S&D'),
        ('gold.brazil_balance_sheet', 'Brazil S&D'),
    ]),
    ('Crop Conditions', [
        ('gold.corn_condition_latest', 'Current corn condition vs 5yr avg'),
        ('gold.soybean_condition_latest', 'Current soybean condition'),
        ('gold.wheat_condition_latest', 'Current wheat condition'),
        ('gold.nass_condition_yoy', 'Year-over-year condition comparison'),
    ]),
    ('CFTC Positioning', [
        ('gold.cftc_sentiment', 'Managed money positioning summary'),
        ('gold.cftc_corn_positioning', 'Corn MM positions'),
        ('gold.cftc_soybean_positioning', 'Soybean MM positions'),
        ('gold.cftc_wheat_positioning', 'Wheat MM positions'),
    ]),
    ('Energy / Biofuels', [
        ('gold.eia_ethanol_weekly', 'Ethanol production + stocks'),
        ('gold.eia_petroleum_weekly', 'Petroleum data'),
        ('gold.eia_prices_daily', 'Daily energy prices'),
        ('gold.emts_monthly_matrix', 'EPA RIN generation by type/tab'),
        ('gold.rin_monthly_trend', 'RIN generation trends'),
    ]),
    ('Trade & Inspections', [
        ('gold.fgis_inspections_monthly_matrix_kbu', 'Monthly inspections by destination (000 bu)'),
        ('gold.fgis_inspections_weekly_matrix_kbu', 'Weekly inspections by destination (000 bu)'),
        ('gold.futures_daily_validated', 'Validated daily futures prices'),
    ]),
    ('Brazil Production', [
        ('gold.brazil_soybean_production', 'Brazil soy by state'),
        ('gold.brazil_corn_production', 'Brazil corn by state'),
        ('gold.brazil_national_production', 'Brazil national totals'),
        ('gold.brazil_crop_summary', 'Latest Brazil crop estimates'),
    ]),
]

for section_name, views in sections:
    doc.add_heading(section_name, level=3)
    t = doc.add_table(rows=len(views) + 1, cols=2, style='Light Grid Accent 1')
    t.rows[0].cells[0].text = 'View'
    t.rows[0].cells[1].text = 'Description'
    for p in t.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in t.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for i, (view, desc) in enumerate(views):
        t.rows[i + 1].cells[0].text = view
        t.rows[i + 1].cells[1].text = desc

# Troubleshooting
doc.add_heading('Troubleshooting', level=2)

issues = [
    ('"Database connection failed" in Excel',
     'Verify ODBC driver is installed (Step 2) and ActiveX reference is checked (Step 3). '
     'Check your internet connection \u2014 the database is on AWS.'),
    ('"no pg_hba.conf entry for host" error',
     'Your IP address needs to be added to the AWS security group. '
     'Contact Tore with your public IP (find it at https://whatismyip.com). '
     'Your ISP may change your IP periodically \u2014 if it stops working, check again.'),
    ('"no encryption" or SSL error',
     'Make sure you have the latest .bas files \u2014 they include sslmode=require in the '
     'connection string. In pgAdmin, set SSL mode to "Require" on the SSL tab.'),
    ('VBA shortcut not working after import',
     'Close and reopen the workbook. Make sure the workbook is saved as .xlsm '
     '(macro-enabled). Check that the ThisWorkbook module has the shortcut assignment code.'),
    ('IP address changes',
     'If you lose access after it was working, your ISP likely rotated your IP. '
     'Check https://whatismyip.com and send the new IP to Tore to update the security group.'),
]
for title, fix in issues:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(fix)

out = 'docs/analyst_database_setup_guide.docx'
doc.save(out)
print(f'Saved: {out}')
